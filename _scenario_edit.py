# -*- coding: utf-8 -*-
"""기획서 docx: 실제 화면 교체 + 캡션 시나리오化 + §2 시나리오 도입부 + §5 문화데이터 증거.
백업본을 원본으로 읽어 캐노니컬로 저장(재실행 안전)."""
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\dev\contest\제출본\2. 기획서_이음_백업_20260621_170252.docx"
DST = r"C:\dev\contest\제출본\2. 기획서_이음.docx"
W = 2.55  # 그리드 이미지 폭(inch)

d = Document(SRC)
part = d.part

# ── 폰트 ──────────────────────────────────────────────
def setf(run, size=9.5, bold=False, color=None):
    run.font.name = "맑은 고딕"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), "맑은 고딕")
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color

def set_para_text(para, text, size, bold):
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    setf(para.add_run(text), size, bold)

def text_paras(cell):
    out = []
    for para in cell.paragraphs:
        if para._p.findall('.//' + qn('w:drawing')):
            continue
        if para.text.strip():
            out.append(para)
    return out

def set_caption(cell, title, desc):
    tps = text_paras(cell)
    if tps:               set_para_text(tps[0], title, 9.5, True)
    if len(tps) > 1:      set_para_text(tps[1], desc, 8.5, True)

def put_image(cell, blob, w=W):
    target = None
    for para in cell.paragraphs:
        if para._p.findall('.//' + qn('w:drawing')):
            target = para; break
    if target is None:
        target = cell.paragraphs[0].insert_paragraph_before()
    for r in list(target.runs):
        r._element.getparent().remove(r._element)
    target.add_run().add_picture(io.BytesIO(blob), width=Inches(w))

# ── 1) loose 이미지 blob 확보(삭제 전) ───────────────────
blob = {}
for rid in ("rId1", "rId2", "rId3", "rId4", "rId5", "rId6"):
    blob[rid] = part.related_parts[rid].blob

tbl = d.tables[1]
sec2 = tbl.rows[5].cells[1]
sec6 = tbl.rows[13].cells[1]
nt2u = sec2.tables[0]   # 사용자 앱 3×2
nt2d = sec2.tables[1]   # 대시보드 2×2
nt6  = sec6.tables[0]   # AI 화면 5×2

# ── 2) §6 실제화면 교체 + 시나리오 캡션 ──────────────────
put_image(nt6.rows[0].cells[0], blob["rId2"])  # 주간주제관리(AI질의생성)
set_caption(nt6.rows[0].cells[0], "① 복지사가 AI로 이번 주 질문을 만든다",
            "'성황당산성'(경남 사천) 주제를 넣자 AI가 노인·청년 맞춤 설문을 자동 생성. 자동·수동(프롬프트 직접) 미리보기 선택.")
put_image(nt6.rows[0].cells[1], blob["rId3"])  # 설문 미리보기 AI초안
set_caption(nt6.rows[0].cells[1], "② AI 초안을 복지사가 검수·협의한다",
            "보기마다 숨은 MBTI 성향 태그까지 생성. 마음에 안 들면 AI와 대화로 다듬어(refine) 발행.")
set_caption(nt6.rows[1].cells[0], "③ AI에 보내는 프롬프트와 받는 양식",  # 유지(프롬프트 양식 mockup)
            "템플릿 자동완성으로 '결정적 프롬프트' 요청 → 구조화 JSON 수령. 키 없이 외부 AI로 수동 처리도 가능.")
put_image(nt6.rows[1].cells[1], blob["rId4"])  # 노인 설문 응답
set_caption(nt6.rows[1].cells[1], "④ 이순자 어르신이 질문을 받는다",
            "선택지를 고르고, 못다 한 말은 '음성으로 말하기'로. 큰 글씨·음성 입력으로 어르신도 쉽게.")
set_caption(nt6.rows[2].cells[0], "⑤ 같은 질문이 청년에게 닿는다",
            "익명·다크 테마로 부담 없이. 두 세대가 같은 문화 이야기로 연결된다.")
set_caption(nt6.rows[2].cells[1], "⑥ 어르신의 음성이 글이 된다",
            "STT로 텍스트화되어 AI와 청년에게 전달된다.")
set_caption(nt6.rows[3].cells[0], "⑦ AI가 공감하며 대화를 잇는다",
            "답변을 분석해 후속 질문으로 대화를 이어가고, 어르신껜 TTS로 읽어준다.")
set_caption(nt6.rows[3].cells[1], "⑧ 모인 이야기로 결과물을 만든다",
            "여러 세대의 답변을 모아 수필·시·그림을 AI가 생성한다.")
put_image(nt6.rows[4].cells[0], blob["rId6"])  # AI 수필
set_caption(nt6.rows[4].cells[0], "⑨ 흩어진 기억이 한 편의 글이 된다",
            "6명의 이야기가 수필 '산성에 남은 이야기'로 재탄생, 음성 낭독으로 어르신께 되돌린다.")
set_caption(nt6.rows[4].cells[1], "⑩ 같은 주제가 그림으로도 남는다",
            "지역문화가 새로운 콘텐츠로 재생산된다.")

# ── 3) §2 사용자 앱 캡션(빈칸 ③ 채우기) ──────────────────
set_caption(nt2u.rows[0].cells[0], "① 로그인", "이메일 인증 후 노인·청년 맞춤 홈으로 분기.")
set_caption(nt2u.rows[0].cells[1], "② 개인정보 보호·이용 동의", "항목별 동의·비식별·복지사 열람 명시.")
import os
elder = os.path.join(r"C:\dev\contest\제출본\img", "03_topic_elder.png")
with open(elder, "rb") as f:
    put_image(nt2u.rows[1].cells[0], f.read())   # 빈칸: 노인 주제카드(mockup)
set_caption(nt2u.rows[1].cells[0], "③ 이번 주 이야기 도착", "매주 지역문화 주제를 큰 글씨·사진으로 제시.")
set_caption(nt2u.rows[1].cells[1], "④ 청년 화면(익명·다크)", "익명 배지로 부담 없이 참여.")
set_caption(nt2u.rows[2].cells[0], "⑤ 참여 통계 보기", "'같은 생각을 한 사람들'과의 연결감.")
set_caption(nt2u.rows[2].cells[1], "⑥ 위기 감지·1393 긴급 연결", "위기 신호 시 자살예방상담 1393·복지사 알림.")

# ── 4) §2 대시보드 캡션(빈칸 채우기: 실제 대시보드) ───────
put_image(nt2d.rows[0].cells[0], blob["rId5"])  # 실제 복지사 대시보드
set_caption(nt2d.rows[0].cells[0], "복지사 대시보드 — 실시간 모니터링",
            "김복지(전북 순창) 화면: 긴급2·주의5·정상18, '긴급' 이순자 어르신 즉시 식별.")
set_caption(nt2d.rows[0].cells[1], "설문 분석 — 응답 통계+누적 MBTI",
            "응답 분포와 누적 성향(MBTI 4축)을 함께.")
set_caption(nt2d.rows[1].cells[0], "멀티 AI 공급자·키 관리",
            "Claude·GPT·Gemini·OpenCode 등록, 실패 시 자동 전환.")
set_caption(nt2d.rows[1].cells[1], "회원·복지사 관리",
            "담당 배정/해제·역할(현장/상위) 관리.")

# ── 5) §2 떠 있는 loose 이미지 6장 제거 ──────────────────
removed = 0
for para in list(sec2.paragraphs):
    if para._p.findall('.//' + qn('w:drawing')):
        para._p.getparent().remove(para._p)
        removed += 1
print("loose paragraphs removed from §2:", removed)

# ── 6) §2 시나리오 도입부 삽입(핵심 흐름 앞) ──────────────
narrative = ("경남 사천의 옛이야기 '성황당산성'이 이번 주 주제로 발행된다. "
             "78세 이순자 어르신은 큰 글씨로 도착한 질문에 옛 기억을 떠올리며 음성으로 답한다. "
             "같은 질문을 받은 한 청년은 익명으로 자신의 생각을 남긴다. AI는 두 세대의 답을 잇고 감정·언어 패턴을 분석한다. "
             "며칠 뒤 접속이 끊기고 감정이 가라앉자 복지사 김복지에게 '긴급' 알림이 뜬다 — 전화 한 통이 닿는다. "
             "그렇게 모인 이야기들은 수필 '산성에 남은 이야기' 한 편으로 다시 태어난다. "
             "문화가 대화를 만들고, 대화가 관계를, 관계가 사회안전망을 만든다.")
anchor = None
for para in sec2.paragraphs:
    if para.text.strip().startswith("■ 핵심 사용자 흐름"):
        anchor = para; break
if anchor is not None:
    h = anchor.insert_paragraph_before()
    setf(h.add_run("■ 사용자 시나리오 — 어느 한 주의 이야기"), 10, True, RGBColor(0xC0, 0x39, 0x2B))
    n = anchor.insert_paragraph_before()
    setf(n.add_run(narrative), 9.5, False)
    print("scenario intro inserted")
else:
    print("WARN: 핵심 흐름 anchor not found")

# ── 7) §5 문화데이터 증거 이미지(후보 101건) ─────────────
c5 = tbl.rows[11].cells[0]
ph = c5.add_paragraph()
setf(ph.add_run("■ 실제 적재된 문화데이터 — 주제 후보 선택 화면"), 9.5, True, RGBColor(0xC0, 0x39, 0x2B))
pi = c5.add_paragraph()
pi.add_run().add_picture(io.BytesIO(blob["rId1"]), width=Inches(3.2))
pc = c5.add_paragraph()
setf(pc.add_run("문화 빅데이터 플랫폼에서 적재한 후보 101건 중 선택(예: 안동 하회 염행당 고택). 이미지/텍스트 합계 189건."), 8.5, False)
print("§5 evidence image added")

try:
    d.save(DST)
    print("SAVED:", DST)
except PermissionError:
    alt = DST.replace(".docx", "_새버전.docx")
    d.save(alt)
    print("LOCKED -> SAVED ALT:", alt)
