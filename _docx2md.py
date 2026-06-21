# -*- coding: utf-8 -*-
"""기존 기획서 docx의 내용을 1:1 그대로 MD로 변환(검토용). 내용 변경 없음, 이미지 추출 포함."""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = r"C:\dev\contest\제출본\2. 기획서_이음.docx"
OUT = r"C:\dev\contest\제출본\기획서_본문.md"
IMGDIR = r"C:\dev\contest\제출본\_md_img"
os.makedirs(IMGDIR, exist_ok=True)
# 기존 추출물 정리
for f in os.listdir(IMGDIR):
    os.remove(os.path.join(IMGDIR, f))

d = Document(SRC)
part = d.part
counter = [0]

def save_blob(blob):
    counter[0] += 1
    name = f"img{counter[0]:02d}.png"
    with open(os.path.join(IMGDIR, name), "wb") as f:
        f.write(blob)
    return f"_md_img/{name}"

def para_images(p_el):
    out = []
    for bl in p_el.findall('.//' + qn('a:blip')):
        rid = bl.get(qn('r:embed'))
        if rid and rid in part.related_parts:
            out.append(part.related_parts[rid].blob)
    return out

def cell_lines(cell):
    """셀 안의 텍스트 단락만 순서대로(이미지 제외)."""
    out = []
    for para in cell.paragraphs:
        if para._p.findall('.//' + qn('w:drawing')):
            continue
        t = para.text.strip()
        if t:
            out.append(t)
    return out

def cell_first_image(cell):
    for para in cell.paragraphs:
        imgs = para_images(para._p)
        if imgs:
            return imgs[0]
    return None

md = []

# ── 상단 제목 ──
md.append("# 제4회 문화체육관광 인공지능·데이터 활용 공모전")
md.append("## — 문화데이터 활용 분야 —\n")
md.append("> ⚠️ 이 파일은 검토용입니다 — 기존 `2. 기획서_이음.docx`의 내용을 그대로 옮긴 것입니다.\n")
md.append("---\n")

t = d.tables[1]

def emit_content_cell(cell):
    """내용 셀: 텍스트 단락 + 표 밖 이미지 + 중첩표(이미지 그리드)를 문서 순서대로."""
    # 본문 요소를 순서대로 순회
    body = cell._tc
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn('w:p'):
            para = Paragraph(child, cell)
            text = para.text.strip()
            imgs = para_images(child)
            if text:
                if text.startswith('■'):
                    md.append(f"\n**{text}**\n")
                elif text.startswith('※'):
                    md.append(f"*{text}*  ")
                elif text.startswith('•') or text.startswith('-'):
                    md.append(f"{text}  ")
                elif text.startswith(('[', '→')):
                    md.append(f"{text}  ")
                else:
                    md.append(f"{text}\n")
            for blob in imgs:
                path = save_blob(blob)
                md.append(f'\n<img src="{path}" width="300">\n')
        elif tag == qn('w:tbl'):
            tbl = Table(child, cell)
            md.append("\n<table>")
            for row in tbl.rows:
                md.append("<tr>")
                for c in row.cells:
                    blob = cell_first_image(c)
                    lines = cell_lines(c)
                    cap = "<br>".join(lines)
                    if blob:
                        path = save_blob(blob)
                        md.append(f'<td align="center"><img src="{path}" width="280"><br>{cap}</td>')
                    else:
                        md.append(f'<td align="center">{cap}</td>')
                md.append("</tr>")
            md.append("</table>\n")

# ── 각 행 처리 ──
for ri, row in enumerate(t.rows):
    cell = row.cells[0]
    lines = cell_lines(cell)
    if ri == 0:   # 공모 부문
        val = " ".join(l for l in cell_lines(row.cells[1] if len(row.cells) > 1 else cell))
        md.append(f"**공모 부문**: {val or (lines[0] if lines else '')}\n")
        continue
    if ri == 1:   # 아이디어 명
        val = " ".join(cell_lines(row.cells[1])) if len(row.cells) > 1 else ""
        md.append(f"**아이디어 명**: {val or (lines[0] if lines else '')}\n")
        continue
    # 섹션 제목 행 (1)~6))
    if lines and lines[0][:2] in ("1)", "2)", "3)", "4)", "5)", "6)"):
        md.append(f"\n---\n\n## {lines[0]}\n")
        continue
    # 내용 행
    emit_content_cell(cell)

open(OUT, "w", encoding="utf-8").write("\n".join(md))
print("WROTE:", OUT)
print("images extracted:", counter[0])
