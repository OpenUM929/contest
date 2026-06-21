# -*- coding: utf-8 -*-
"""이음 공모전 신청서/기획서 자동 작성 — 원본 양식 보존, 사본에 기입."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONT = "맑은 고딕"

def _p_text(el):
    return ''.join(t.text or '' for t in el.findall('.//' + qn('w:t')))

def trim_to_form(doc, start_marker, end_marker):
    """start_marker 헤딩 이전, end_marker 헤딩 이후를 모두 제거(최종 sectPr는 보존)."""
    body = doc.element.body
    start = end = None
    for el in list(body):
        if el.tag == qn('w:p'):
            txt = _p_text(el)
            if start is None and start_marker in txt:
                start = el
            elif start is not None and end is None and end_marker in txt:
                end = el
    if start is None:
        print("WARN: start marker not found")
        return
    # start 이전 제거
    removing = True
    for el in list(body):
        if el is start:
            break
        body.remove(el)
    # end 이후 제거(sectPr 제외)
    if end is not None:
        removing = False
        for el in list(body):
            if el is end:
                removing = True
            if removing and el.tag != qn('w:sectPr'):
                body.remove(el)

def set_font(run, name=FONT, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    if color is not None:
        run.font.color.rgb = color

def set_cell(cell, text, bold=False, size=10, color=None):
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    set_font(run, size=size, bold=bold, color=color)

def add_line(cell, text="", bold=False, size=9.5, bullet=False, indent=False, color=None):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    if indent:
        pf.left_indent = Pt(12)
    run = p.add_run(("• " if bullet else "") + text)
    set_font(run, size=size, bold=bold, color=color)
    return p

NAVY = RGBColor(0x1A, 0x1A, 0x2E)

# ============================================================
# 1) 참가신청서
# ============================================================
def fill_application(src, dst):
    doc = Document(src)
    t = doc.tables[1]
    # 공모 분야: '문화데이터 활용' = R1 cells[5] (merged 5-6)
    set_cell(t.rows[1].cells[5], "☑ 문화데이터 활용", bold=True)
    # 공모 부문: '아이디어' = R2 cells[6]
    set_cell(t.rows[2].cells[6], "☑ 아이디어", bold=True)
    # 성명: 팀 = R3 cells[6]
    set_cell(t.rows[3].cells[6], "☑ 팀(이음)   (인원: 총 [[O]]명)", bold=True)
    # 팀장 이메일 (R5, cells[8]) — 알려진 대표 이메일
    set_cell(t.rows[5].cells[8], "feelmydream80@gmail.com")
    # 출품작명 (R10, merged from cells[1])
    set_cell(t.rows[10].cells[1],
             "이음(以音) — 지역문화 데이터로 세대를 잇는 AI 사회안전망 플랫폼", bold=True)
    # 출품작 개요 (R11, merged from cells[1])
    c = t.rows[11].cells[1]
    c.text = ''
    set_font(c.paragraphs[0].add_run(
        "문화 빅데이터 플랫폼의 지역문화 멀티미디어 데이터(설화·지명유래·옛 사진)를 매주 "
        "'세대공감 대화 주제'로 발행하여, 은둔 청년과 독거노인이 같은 문화 기억을 매개로 마음을 "
        "나누도록 잇는 AI 플랫폼이다. AI가 대화 속 감정·심리지표·성향을 분석해 고독사·고립 위기를 "
        "조기 감지·복지사에게 연결하며, 모인 대화는 한 편의 수필로 재탄생해 세대를 잇는다."), size=10)
    doc.save(dst)
    print("saved:", dst)

# ============================================================
# 2) 기획서 [붙임 2-4] = TABLE[7]
# ============================================================
def fill_plan(src, dst):
    doc = Document(src)
    t = doc.tables[7]

    # R0 공모 부문 / R1 아이디어 명
    set_cell(t.rows[0].cells[1], "☑ 아이디어 부문", bold=True)
    set_cell(t.rows[1].cells[1],
             "이음(以音) — 지역문화 데이터로 잇는 세대공감 AI 사회안전망", bold=True, size=11)

    # ---- 1) 아이디어 소개(요약): 답변 = R3 ----
    c = t.rows[3].cells[-1]
    add_line(c, "")
    add_line(c,
        "이음(以音)은 '지역문화 데이터'를 매개로 은둔 청년과 독거노인을 잇는 AI 사회안전망이다. "
        "매주 문화 빅데이터 플랫폼의 지역 설화·지명유래·옛 사진을 '세대공감 대화 주제'로 발행하면, "
        "두 세대가 같은 문화 기억을 떠올리며 대화·설문에 참여한다. AI는 그 대화에서 감정·인지·성향(MBTI) "
        "지표를 분석해 고독사·고립 위기를 조기 감지하고 복지사에게 연결하며, 모인 이야기를 한 편의 수필로 "
        "엮어 세대를 잇는다. 청년의 사회 단절과 노인의 고독사라는 두 사회문제를, 문화가 가진 '공감의 힘'으로 "
        "동시에 푸는 것이 목표다.", size=10)

    # ---- 2) 상세 설명: 답변 = R5 ----
    c = t.rows[5].cells[-1]
    add_line(c, "")
    add_line(c, "■ 문제의식", bold=True, color=NAVY)
    add_line(c, "노인 고독사 — 1인 가구 노인 급증으로 관계 단절이 고독사로 이어지나, 안부 확인은 사람 발품에 의존해 사각지대가 넓다.", bullet=True)
    add_line(c, "청년 고립·은둔 — 사회와 단절된 은둔형 외톨이 청년이 늘지만, 이들을 끌어낼 '역할·동기' 설계가 부족하다.", bullet=True)
    add_line(c, "세대 단절 — 서로의 자원이 될 수 있는 두 세대가 만날 접점이 없다.", bullet=True)
    add_line(c, "→ 이음은 셋을 '지역문화 데이터'라는 하나의 매개로 동시에 연결한다.", indent=True)
    add_line(c, "")
    add_line(c, "■ 핵심 사용자 흐름 (7단계 파이프라인)", bold=True, color=NAVY)
    add_line(c, "[A] 주제 발행 → [B] 세대 대화 → [C] 심리 분석 → [D] 수필 생성 → [E] 안전 감지 → [F] 복지사 연결 → [G] 데이터 보안(횡단)", indent=True)
    add_line(c, "")
    add_line(c, "■ 주요 기능 (전체 기능 카탈로그)", bold=True, color=NAVY)
    add_line(c, "[문화데이터·주제] 공공 문화데이터(문화 빅데이터 플랫폼·국가기록원·민속아카이브) 자동 수집·검증·캐싱 → 지역별 맞춤 주제 → 매주 자동 발행(스케줄러)", bullet=True)
    add_line(c, "[AI 설문 생성] 선택·서술·혼합형 질문 자동 생성, 보기별 MBTI 성향 태깅, 자동·수동 미리보기, 복지사 검수·AI 협의(refine) 후 발행", bullet=True)
    add_line(c, "[세대 대화] 노인 음성(STT→AI→TTS)·청년 텍스트·음원 재생, 글자 크기 등 접근성, AI 공감 응답으로 대화 지속", bullet=True)
    add_line(c, "[심리·성향 분석] 감정 분석, 어휘 다양성(TTR)·반복표현(치매 조기 신호), 누적 MBTI/성향 집계", bullet=True)
    add_line(c, "[안전망] 위기 키워드 감지→1393 연결, 3일 접속 단절·감정 악화 자동 점검(매일 09:00), 복지사 이메일 알림", bullet=True)
    add_line(c, "[복지사 대시보드] 위험도 점수 등급(긴급/주의/정상), 주제 확인 현황, 개입(전화·방문·상담·알림해결) 이력, 설문 분석·응답 조회, 회원·복지사 관리, 멀티 AI 키 관리", bullet=True)
    add_line(c, "[결과물] 모인 대화 → 수필·시·소설·이미지 생성, 기여자 추적, 음성 낭독 전달", bullet=True)
    add_line(c, "[보안·신뢰] 개인정보 항목별 동의, 대화 비식별·암호화, 복지사 검수(Human-in-the-loop)", bullet=True)
    add_line(c, "")
    add_screens(c, SCREENS_USER, label="■ 사용자 앱 화면 (노인·청년 / React Native)")
    add_line(c, "")
    add_screens(c, SCREENS_ADMIN, label="■ 복지사 대시보드 화면 (웹 / React)")

    # ---- 3) 독창성: 답변 = R7 ----
    c = t.rows[7].cells[-1]
    add_line(c, "")
    add_line(c, "1. 문화데이터를 '소비 콘텐츠'가 아니라 '관계·돌봄의 촉매'로 재정의", bold=True)
    add_line(c, "기존 문화데이터 서비스는 추천·전시·관광 등 '보여주기'에 머문다. 이음은 같은 데이터를 세대 간 대화를 여는 열쇠이자 심리 상태를 읽는 진단 도구로 사용한다.", indent=True)
    add_line(c, "2. 두 사회문제(청년 고립 + 노인 고독사)를 하나의 구조로 동시 해결", bold=True)
    add_line(c, "노인은 말벗을, 청년은 사회적 역할을 얻는 상호 호혜 설계. 한쪽을 돕는 행위가 다른 쪽의 치유가 된다.", indent=True)
    add_line(c, "3. '향수 → 자연스러운 자기노출 → 심리 진단'의 비침습적 모니터링", bold=True)
    add_line(c, "직접적 우울·외로움 설문 대신, 옛 기억을 떠올리는 문화 대화 속에서 감정·인지·성향을 자연스럽게 수집해 거부감을 낮춘다.", indent=True)
    add_line(c, "4. 대화의 수필화로 '데이터가 다시 문화가 되는' 선순환", bold=True)
    add_line(c, "문화데이터에서 출발한 대화가 새로운 창작물(수필)로 환원되어 다시 세대를 잇는다.", indent=True)
    add_line(c, "")
    add_line(c, "■ 실제 데이터 기반 — 구상이 아닌 실증", bold=True, color=NAVY)
    add_line(c, "현재 DB에 전국 지역문화 데이터 189건(이미지 101·텍스트 88)이 실제 적재되어, 아래와 같은 실존 주제로 설문이 생성된다.", bullet=True)
    add_line(c, "예: '강천용소와 노총각'(전북 순창군·지역설화) / '황골엿과 황골 엿술'(강원 원주·향토음식) / '성북동 이종석 별장'(서울·한국의 가옥) / '주왕산지'(경북 청송·기록문화)", indent=True)
    add_line(c, "추출된 TOP10 키워드(역사/인물 25 · 전통가옥 13 · 생활/문화 11 · 불교/문화유산 10 · 이야기/설화 9 …)로 복지사가 지역·테마를 골라 맞춤 발행한다. 같은 데이터가 세대공감 대화와 성향 분석으로 직결되는 점이 독창적이다.", bullet=True)

    # ---- 4) 발전 가능성 및 기대효과: 답변 = R9 ----
    c = t.rows[9].cells[-1]
    add_line(c, "")
    add_line(c, "■ 구현·사업화 실현 가능성 — 이미 작동하는 MVP로 검증 완료", bold=True, color=NAVY)
    add_line(c, "본 아이디어는 구상에 그치지 않고, 실제 동작하는 프로토타입(MVP)으로 핵심 파이프라인을 구현·검증했다.", indent=True)
    add_line(c, "백엔드(FastAPI) — 공공 문화데이터 API 연동, AI 주제·질문 자동 생성, 감정·위기 감지, 수필 생성, 복지사 알림(이메일/스케줄러) 동작", bullet=True)
    add_line(c, "복지사 대시보드(React) — 주제 발행·설문 관리·대상자 심리지표·위험도 모니터링 화면 구현", bullet=True)
    add_line(c, "모바일 앱 — 노인 음성 대화·청년 텍스트 대화, 위기 시 1393 연결 팝업 구현", bullet=True)
    add_line(c, "→ 사업화 시 '개념 검증' 단계를 이미 통과한 상태이며, 지자체·복지관 실증(파일럿)으로 곧바로 확장 가능하다.", indent=True)
    add_line(c, "")
    add_line(c, "■ 사회적 가치 (정성적 기대효과)", bold=True, color=NAVY)
    add_line(c, "고독사 예방 — 비대면 안부·심리 모니터링으로 위기를 조기 포착, 복지 사각지대 축소", bullet=True)
    add_line(c, "청년 사회 재연결 — '누군가에게 필요한 존재'라는 역할 경험으로 은둔 탈출 동기 부여", bullet=True)
    add_line(c, "세대 통합 — 문화 기억을 공유하며 세대 간 이해와 공감 회복", bullet=True)
    add_line(c, "지역문화 가치 재발견 — 잊혀가던 지역 설화·지명유래가 일상 대화 속에서 다시 살아남", bullet=True)
    add_line(c, "사회적 포용 — 디지털 약자(노인 음성 UI)·다문화·비수도권 지역문화까지 포괄하는 포용적·다양성 설계", bullet=True)
    add_line(c, "")
    add_line(c, "■ 정량적 기대효과 (실증 단계 KPI)", bold=True, color=NAVY)
    add_line(c, "복지사 1인당 모니터링 가능 대상자 수 증가(수기 점검 → 대시보드 자동 집계), 위기 감지 소요시간 단축(정기 방문 → 실시간), 주당 대화·수필 생성 건수 등 참여 지표 측정", bullet=True)
    add_line(c, "")
    add_line(c, "■ 확장성·파급효과", bold=True, color=NAVY)
    add_line(c, "음원(국악)·영상 등 멀티미디어 문화데이터 및 다국어(다문화가정)로 확대 / 시니어 케어·디지털 헬스케어·지역문화 콘텐츠 산업과 결합 / 지자체·경로당·치매안심센터 연계 공공 돌봄 인프라화", bullet=True)
    add_line(c, "")
    add_line(c, "■ 확산 경로 (사업화)", bold=True, color=NAVY)
    add_line(c, "지자체·복지관과 B2G 실증(파일럿) → 효과 검증 후 광역단위 확대 → 경로당·치매안심센터·청년지원센터 연계로 전국 확산", bullet=True)
    add_line(c, "")
    add_line(c, "■ ESG 실현 가능성", bold=True, color=NAVY)
    add_line(c, "사회(S) — 청년 사회참여 역할 창출·복지 인력 보조로 일자리 기여, 디지털 격차·다문화·지역 균형 등 포용성 제고", bullet=True)
    add_line(c, "환경(E) — 비대면 안부·모니터링으로 대면 방문 부담을 줄여 이동·탄소 저감에 기여", bullet=True)
    add_line(c, "지배구조(G) — 대화 데이터 비식별·암호화 및 복지사 검수(Human-in-the-loop)로 윤리적 AI 운영", bullet=True)

    # ---- 5) 문화데이터 활용: 답변 = R11 (빈 셀) ----
    c = t.rows[11].cells[-1]
    add_line(c, "■ 활용 데이터", bold=True, color=NAVY)
    add_line(c, "[필수] 지역문화 멀티미디어 데이터(설화·전설·지명유래·옛 사진) — 제공: 한국문화정보원 / 출처: 문화 빅데이터 플랫폼(bigdata-culture.kr, 문화체육관광부 데이터플랫폼)", bullet=True)
    add_line(c, "데이터셋 상세 URL: [[문화 빅데이터 플랫폼 데이터셋 페이지 URL 기입]]", indent=True)
    add_line(c, "[선택·융복합] 국가기록원 나라기록물정보(공공데이터포털) / 국립민속박물관 민속아카이브(KCISA 문화공공데이터광장)", bullet=True)
    add_line(c, "")
    add_line(c, "■ 실제 활용 규모 (MVP 적재 기준)", bold=True, color=NAVY)
    add_line(c, "분기별 원천 데이터(2022.12~2023.12) 5종을 수집·병합 → 이미지 주제 101건 + 텍스트 주제 88건, 총 189건의 지역문화 콘텐츠를 실제 적재하고 TOP10 키워드를 추출해 운용", bullet=True)
    add_line(c, "")
    add_line(c, "■ 활용 방식 (전처리·가공·연계)", bold=True, color=NAVY)
    add_line(c, "수집 — 문화 빅데이터 플랫폼에서 지역문화 데이터(CSV)를 내려받고, 국가기록원·민속아카이브 API를 실시간 연동", bullet=True)
    add_line(c, "전처리 — 이미지 URL 유효성 검증, 설화 텍스트 정제, 미디어 로컬 캐싱", bullet=True)
    add_line(c, "분할·구조화 — 이미지 유무로 image/text 주제 분류, 핵심 키워드(TOP10) 추출 → 복지사 대시보드 필터로 제공", bullet=True)
    add_line(c, "AI 가공 — 데이터의 제목·설명·키워드를 AI에 입력해 연령대별 세대공감 대화 질문으로 변환", bullet=True)
    add_line(c, "")
    add_line(c, "■ 데이터의 역할 — 문화데이터는 ①대화를 여는 '주제', ②심리·성향을 읽는 '진단 매개', ③수필 창작의 '소재'라는 3중 역할을 수행한다. 즉, 단순 표시 대상이 아니라 세대를 잇고 위기를 감지하는 서비스의 작동 원리 그 자체다.", bold=True)

    # ---- 6) AI 기술 활용: 답변 = R13 ----
    c = t.rows[13].cells[-1]
    add_line(c, "")
    add_line(c, "■ 적용 영역별 AI 기술", bold=True, color=NAVY)
    add_line(c, "주제·질문 생성 — 멀티프로바이더 LLM(Claude·GPT·Gemini)+프롬프트 엔지니어링: 문화데이터→연령 맞춤 세대공감 질문(선택/서술/혼합형) 자동 생성", bullet=True)
    add_line(c, "세대 대화 — LLM 대화 엔진+STT/TTS: 노인 음성↔텍스트 변환, 연령·역할별 페르소나 응답", bullet=True)
    add_line(c, "심리 분석 — 감정분류 모델(KR-FinBert)+언어패턴 분석(TTR·n-gram)+성향(MBTI) 태깅: 대화에서 감정·인지·성향 지표 추출", bullet=True)
    add_line(c, "위기 감지 — 위기 키워드 탐지+규칙·패턴 분류: 자해·고립 신호를 high/medium/low로 분류해 대응 분기", bullet=True)
    add_line(c, "수필 생성 — 생성형 LLM: 다수의 대화를 한 편의 수필로 자동 창작", bullet=True)
    add_line(c, "")
    add_line(c, "■ AI 활용 특징", bold=True, color=NAVY)
    add_line(c, "다중 AI 공급자 자동 전환(Fallback) — 한 AI가 실패해도 다른 공급자로 자동 전환해 서비스 연속성 확보", bullet=True)
    add_line(c, "결정적 프롬프트+견고한 파싱 — 코드블록·산문이 섞인 AI 응답도 구조화 JSON으로 안정 추출", bullet=True)
    add_line(c, "사람 개입(Human-in-the-loop) — 복지사가 AI 생성 주제·질문을 검토·수정·발행하는 자동/수동 미리보기 체계로 품질·안전성 보장", bullet=True)
    add_line(c, "")
    add_line(c, "■ 프롬프트 요청 방법과 수령 양식", bold=True, color=NAVY)
    add_line(c, "요청 — 질문 유형별 프롬프트 템플릿에 주제·대상·미디어 유형을 자동 대입(template.format)해 '결정적 프롬프트'를 만들어 LLM에 전달. AI 연결이 어려운 환경을 위해, 같은 프롬프트 원문을 화면에 출력하고 복지사가 외부 AI 답변을 붙여넣는 '수동 미리보기'도 지원한다.", bullet=True)
    add_line(c, "수령 — schema_version·question_type·questions[](선택지별 MBTI 성향 태그 포함)로 구성된 구조화 JSON(QuestionSet) 양식으로 받는다. 코드블록·설명이 섞인 응답도 견고하게 파싱해 화면 미리보기로 변환한다.", bullet=True)
    add_line(c, "")
    add_screens(c, SCREENS_AI, label="■ 구현 화면 — 위 AI 기술이 실제 작동하는 모습")

    # [붙임 2-4]만 남기고 나머지 폼 제거 → 열면 바로 우리 내용이 보이도록
    trim_to_form(doc, "[붙임 2-4]", "[붙임 2-5]")

    try:
        doc.save(dst)
    except PermissionError:
        alt = dst.replace(".docx", "_새버전.docx")
        doc.save(alt)
        print("WARN: 원본이 열려 있어 대체 저장 ->", alt)
        return
    print("saved:", dst)

IMGDIR = r"C:\dev\contest\제출본\img"

# 2) 아이디어 상세 설명 — 사용자 앱(모바일) 화면
SCREENS_USER = [
    ("01_login.png", "로그인 (공통)", "이메일 인증 후 사용자 유형(노인/청년)별 맞춤 홈으로 분기."),
    ("02_consent.png", "개인정보 보호·이용 동의", "처리방침·약관·AI분석을 항목별 동의. 비식별·복지사 열람 명시로 신뢰 확보."),
    ("03_topic_elder.png", "세대공감 주제 카드(노인)", "매주 지역문화 '이번 주 이야기'를 큰 글씨·사진으로 제시, 진입 시 음성(TTS) 안내."),
    ("08_youth_home.png", "청년 화면(익명·다크)", "익명 배지로 부담 없이 참여, 음원 주제는 바로 재생."),
    ("09_stats.png", "참여 통계 보기", "응답 분포를 막대그래프로 보여 '같은 생각을 한 사람들'과의 연결감 형성."),
    ("07_crisis.png", "위기 감지·1393 긴급 연결", "위기 신호 감지 시 자살예방상담 1393 연결·복지사 알림 팝업."),
]

# 2) 아이디어 상세 설명 — 복지사 대시보드(웹) 화면
SCREENS_ADMIN = [
    ("15_welfare_dashboard.png", "복지사 대시보드 — 실시간 모니터링", "담당자별 위험도(긴급/주의/정상) 점수 등급, 주제 확인 현황, 전화·방문·상담·알림해결 등 개입 이력 관리."),
    ("16_survey_analytics.png", "설문 분석 — 응답 통계+누적 MBTI", "응답률·선택 분포와 함께 누적 성향(MBTI 4축)을 시각화해 대상자 이해를 돕는다."),
    ("17_settings_apikey.png", "멀티 AI 공급자·키 관리", "Claude·GPT·Gemini·OpenCode 등록, 실패 시 자동 전환(Fallback)으로 서비스 연속성 확보."),
    ("18_admin_members.png", "회원·복지사 관리", "복지사 등록, 담당 노인·청년 배정/해제, 역할(현장/상위 관리자) 관리."),
]

# 6) AI 기술 활용에 들어갈 AI 파이프라인 화면
SCREENS_AI = [
    ("11_welfare_topic.png", "복지사 — AI 질의 생성", "주간 주제 관리에서 지역문화 주제를 넣으면 AI가 연령 맞춤 설문을 자동 생성(자동/수동 미리보기)."),
    ("19_refine_chat.png", "복지사 — AI와 설문 협의(refine)", "복지사가 AI와 대화하며 보기·문구를 다듬어 발행 전 완성도를 높인다."),
    ("12_prompt_format.png", "프롬프트 요청·수령 양식", "템플릿 자동완성 '결정적 프롬프트' 요청 → 구조화 JSON(QuestionSet) 수령."),
    ("04_survey_choice.png", "노인 — 질의서(선택형)", "AI가 만든 향수 자극 선택형. 보기마다 숨은 MBTI 태그로 성향 데이터 수집."),
    ("14_youth_survey.png", "청년 — 질의서(선택형)", "동일 설문을 청년 감성 다크 테마로 제공."),
    ("05_survey_voice.png", "서술형 음성 설문 답변", "버튼을 누르고 말하면 STT로 텍스트화. 어르신도 음성으로 서술 응답."),
    ("06_ai_response.png", "AI 공감 응답", "응답 분석 후 후속 질문으로 대화 유도(노인은 TTS 자동 재생)."),
    ("13_deliverable_gen.png", "결과물 생성 페이지", "모인 대화로 수필·시·이미지를 AI가 생성."),
    ("10_essay.png", "결과물 — 수필", "여러 세대의 이야기가 한 편의 글로 재탄생, 음성 낭독 지원."),
    ("20_image_result.png", "결과물 — AI 그림", "같은 주제를 삽화로도 생성해 수필·시와 함께 전달."),
]

def add_screens(cell, items, cols=2, label="■ 주요 화면"):
    import math
    p = cell.add_paragraph(); r = p.add_run(label); set_font(r, size=10.5, bold=True, color=NAVY)
    nrows = math.ceil(len(items) / cols)
    tbl = cell.add_table(rows=nrows, cols=cols)
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for idx, (fn, cap, desc) in enumerate(items):
        cl = tbl.rows[idx // cols].cells[idx % cols]
        cl.text = ""
        path = os.path.join(IMGDIR, fn)
        if os.path.exists(path):
            cl.paragraphs[0].add_run().add_picture(path, width=Inches(2.55))
        cp = cl.add_paragraph(); cr = cp.add_run(cap); set_font(cr, size=9.5, bold=True, color=NAVY)
        dp = cl.add_paragraph(); dr = dp.add_run(desc); set_font(dr, size=8.5)

if __name__ == "__main__":
    base = r"C:\Users\Park\Downloads\공모전 신청서 최종\신청서 양식(word)"
    out = r"C:\dev\contest\제출본"
    os.makedirs(out, exist_ok=True)
    fill_application(
        os.path.join(base, "1. 참가신청서(파일명_1. 참가신청서_팀명).docx"),
        os.path.join(out, "1. 참가신청서_이음.docx"))
    fill_plan(
        os.path.join(base, "2. 기획서(파일명_2. 기획서_팀명)_최종.docx"),
        os.path.join(out, "2. 기획서_이음.docx"))
    print("DONE")
